import ollama
import streamlit as st
import chardet
import PyPDF2
import docx
import io
import time
import logging
from typing import List, Optional, Dict, Tuple
from dataclasses import dataclass
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OllamaEmbeddings  # Updated import
from langchain_community.vectorstores import FAISS  #

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """Metadata for processed documents"""
    filename: str
    chunk_count: int
    total_tokens: int
    processing_time: float

class DocumentProcessor:
    """Handles document processing and metadata tracking"""
    
    def __init__(self):
        self.supported_extensions = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'txt': self._process_text,
            'md': self._process_text,
            'csv': self._process_text
        }
    
    def _process_pdf(self, file) -> str:
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            return '\n'.join(
                page.extract_text().strip()
                for page in pdf_reader.pages
                if page.extract_text().strip()
            )
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise ValueError(f"Failed to process PDF: {str(e)}")

    def _process_docx(self, file) -> str:
        try:
            doc = docx.Document(file)
            # Include headers, paragraphs and tables
            content = []
            
            # Process headers
            for section in doc.sections:
                header = section.header
                if header:
                    content.extend(paragraph.text for paragraph in header.paragraphs)
            
            # Process main content
            content.extend(paragraph.text for paragraph in doc.paragraphs)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    content.extend(cell.text for cell in row.cells)
            
            return '\n'.join(filter(None, content))
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise ValueError(f"Failed to process DOCX: {str(e)}")

    def _process_text(self, file) -> str:
        try:
            raw_data = file.getvalue()
            result = chardet.detect(raw_data)
            encodings = [result['encoding'], 'utf-8', 'latin-1', 'ascii']
            
            for encoding in encodings:
                try:
                    if encoding:
                        return raw_data.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Unable to decode file with any supported encoding")
        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}")
            raise ValueError(f"Failed to process text file: {str(e)}")

    def process_file(self, file) -> Tuple[str, DocumentMetadata]:
        """Process a file and return its content with metadata"""
        start_time = time.time()
        
        filename = file.name
        file_ext = filename.split('.')[-1].lower()
        
        if file_ext not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        processor = self.supported_extensions[file_ext]
        content = processor(file)
        
        # Calculate basic metrics
        metadata = DocumentMetadata(
            filename=filename,
            chunk_count=len(content.split('\n')),
            total_tokens=len(content.split()),
            processing_time=time.time() - start_time
        )
        
        return content, metadata

class EnhancedRAGApplication:
    def __init__(
        self,
        model_name: str = "llama3.2",
        embedding_model: str = "llama3.2",
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        temperature: float = 0.7,
        max_tokens: int = 2048
    ):
        """Initialize the enhanced RAG application"""
        self.model_name = model_name
        self.temperature = temperature
        self.max_tokens = max_tokens
        
        # Initialize components
        self.embeddings = OllamaEmbeddings(
            model=embedding_model,
            base_url="http://localhost:11434"
        )
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        self.vector_store = None
        self.doc_processor = DocumentProcessor()
        self.document_metadata: Dict[str, DocumentMetadata] = {}

    def retrieve_relevant_docs(
        self,
        query: str,
        top_k: int = 3,
        similarity_threshold: float = 0.0
    ) -> List[str]:
        """
        Retrieve most relevant document chunks for a given query.
        
        Args:
            query (str): User query
            top_k (int): Number of top relevant documents to retrieve
            similarity_threshold (float): Minimum similarity score threshold
        
        Returns:
            List[str]: Most relevant document chunks
        """
        if self.vector_store is None:
            raise ValueError("No documents have been loaded. Call load_documents first.")
        
        try:
            # Get documents with scores
            docs_and_scores = self.vector_store.similarity_search_with_score(
                query=query,
                k=top_k
            )
            
            # Filter by similarity threshold and extract content
            relevant_docs = [
                doc.page_content
                for doc, score in docs_and_scores
                if score >= similarity_threshold
            ]
            
            logger.info(f"Retrieved {len(relevant_docs)} relevant documents for query")
            
            return relevant_docs
            
        except Exception as e:
            logger.error(f"Error retrieving relevant documents: {str(e)}")
            raise
    
    def preprocess_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = ' '.join(text.split())
        # Basic text normalization
        text = text.replace('\t', ' ').strip()
        return text

    def load_documents(self, documents: List[str], metadata: Optional[Dict[str, DocumentMetadata]] = None) -> None:
        """Process and index documents with metadata tracking"""
        try:
            doc_chunks = []
            
            for doc in documents:
                cleaned_doc = self.preprocess_text(doc)
                if cleaned_doc:
                    chunks = self.text_splitter.split_text(cleaned_doc)
                    doc_chunks.extend(chunks)
            
            if not doc_chunks:
                raise ValueError("No valid document content found after processing.")
            
            self.vector_store = FAISS.from_texts(
                texts=doc_chunks,
                embedding=self.embeddings
            )
            
            if metadata:
                self.document_metadata.update(metadata)
            
            logger.info(f"Successfully loaded {len(doc_chunks)} chunks from {len(documents)} documents")
            
        except Exception as e:
            logger.error(f"Error loading documents: {str(e)}")
            raise

    def generate_response(
        self,
        query: str,
        context: Optional[List[str]] = None,
        max_context_length: int = 2000
    ) -> Dict:
        """Generate an enhanced response with metadata"""
        start_time = time.time()
        
        try:
            if context is None:
                context = self.retrieve_relevant_docs(query)
            
            # Truncate context if too long
            combined_context = ' '.join(context)
            if len(combined_context) > max_context_length:
                combined_context = combined_context[:max_context_length] + "..."
            
            prompt = self._construct_prompt(query, combined_context)
            
            response = ollama.chat(
                model=self.model_name,
                messages=[
                    {
                        'role': 'system',
                        'content': 'You are a helpful assistant that provides detailed, accurate answers based on the given context. If the context doesn\'t contain enough information to fully answer the question, acknowledge this and provide the best possible answer with the available information.'
                    },
                    {
                        'role': 'user',
                        'content': prompt
                    }
                ],
                options={
                    'temperature': self.temperature,
                    'max_tokens': self.max_tokens
                }
            )
            
            return {
                'content': response['message']['content'],
                'metadata': {
                    'processing_time': time.time() - start_time,
                    'context_length': len(combined_context),
                    'query_length': len(query)
                }
            }
            
        except Exception as e:
            logger.error(f"Error generating response: {str(e)}")
            raise

    def _construct_prompt(self, query: str, context: str) -> str:
        """Construct an enhanced prompt template"""
        return f"""
        Context Information:
        {context}

        Question: {query}

        Please provide a comprehensive answer based on the context above. Consider:
        1. Direct relevance to the question
        2. Accuracy of information
        3. Completeness of response
        4. Clarity and coherence

        Answer:
        """

def create_streamlit_ui():
    """Create an enhanced Streamlit UI"""
    st.set_page_config(
        page_title="Deepseek AI Document Q&A",
        page_icon="🤖",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Custom CSS for better styling
    st.markdown("""
        <style>
        .stButton>button {
            width: 100%;
            margin-top: 1rem;
        }
        .success-message {
            padding: 1rem;
            background-color: #d4edda;
            border-color: #c3e6cb;
            color: #155724;
            border-radius: 0.25rem;
            margin-bottom: 1rem;
        }
        </style>
    """, unsafe_allow_html=True)
    
    # Initialize session state
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
        
    if 'rag_app' not in st.session_state:
        st.session_state.rag_app = None
    
    return st.session_state

def main():
    session_state = create_streamlit_ui()
    
    st.title("🔍 Deepseek & LLAMA RAG Application")
    st.markdown("""
    ### Advanced Document Analysis and Question Answering
    Upload documents or paste text to get AI-powered insights and answers.
    """)
    
    # Initialize RAG application with configurable parameters
    with st.sidebar:
        st.header("⚙️ Configuration")
        model_name = st.selectbox(
                                "Model", 
                                ["llama3.2", "deepseek-r1:7b", "llama2-uncensored", "mistral"],  # Added deepseek
                                index=0
                            )
        temperature = st.slider("Temperature", 0.0, 1.0, 0.7)
        chunk_size = st.number_input("Chunk Size", 100, 1000, 500)
        
        st.header("📁 Document Input")
        uploaded_files = st.file_uploader(
            "Upload Documents",
            type=['txt', 'pdf', 'docx'],
            accept_multiple_files=True,
            help="Supported formats: PDF, DOCX, TXT"
        )
        
        manual_docs = st.text_area(
            "Or Paste Text",
            height=200,
            help="Enter text directly. Each paragraph will be processed separately."
        )
        
        # Create or update RAG app when configuration changes
        if (session_state.rag_app is None or 
            session_state.rag_app.model_name != model_name or 
            session_state.rag_app.temperature != temperature):
            
            session_state.rag_app = EnhancedRAGApplication(
                model_name=model_name,
                temperature=temperature,
                chunk_size=chunk_size
            )
    
    # Main content area
    col1, col2 = st.columns([2, 1])
    
    with col1:
        query = st.text_input("🔍 Enter your question:")
        process_button = st.button("Generate Response")
    
    with col2:
        st.markdown("### 📊 Document Stats")
        if uploaded_files:
            for file in uploaded_files:
                st.info(f"📄 {file.name}")
    
    if process_button and (uploaded_files or manual_docs) and query:
        try:
            with st.spinner('Processing documents and generating response...'):
                # Process documents
                documents = []
                metadata = {}
                
                # Handle uploaded files
                if uploaded_files:
                    for file in uploaded_files:
                        content, doc_metadata = session_state.rag_app.doc_processor.process_file(file)
                        documents.append(content)
                        metadata[file.name] = doc_metadata
                
                # Handle manual input
                if manual_docs:
                    content = manual_docs.strip()
                    if content:
                        documents.append(content)
                        metadata['manual_input'] = DocumentMetadata(
                            filename='manual_input',
                            chunk_count=len(content.split('\n')),
                            total_tokens=len(content.split()),
                            processing_time=0.0
                        )
                
                # Load documents and generate response
                session_state.rag_app.load_documents(documents, metadata)
                response_data = session_state.rag_app.generate_response(query)
                
                # Display results in an organized way
                st.markdown("### 📝 Response")
                st.write(response_data['content'])
                
                # Display metadata
                with st.expander("🔍 Response Details"):
                    st.json(response_data['metadata'])
                
                # Add to chat history
                session_state.chat_history.append({
                    'question': query,
                    'answer': response_data['content'],
                    'metadata': response_data['metadata']
                })
                
        except Exception as e:
            st.error(f"Error: {str(e)}")
            logger.error(f"Application error: {str(e)}", exc_info=True)
    
    # Display chat history
    if session_state.chat_history:
        st.markdown("### 💬 Chat History")
        for i, chat in enumerate(session_state.chat_history):
            with st.expander(f"Q: {chat['question'][:50]}..."):
                st.write("Question:", chat['question'])
                st.write("Answer:", chat['answer'])
                st.write("Metadata:", chat['metadata'])

if __name__ == "__main__":
    main()