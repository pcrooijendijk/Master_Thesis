import torch
import logging
import docx
import PyPDF2
import chardet
import time
import os
import re
import pickle
from langchain_community.vectorstores import FAISS   
from typing import Tuple, List, Optional, Dict
from dataclasses import dataclass

from transformers import AutoTokenizer, AutoModelForCausalLM, AutoConfig, GenerationConfig
from langchain_text_splitters import CharacterTextSplitter, RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document
from langchain_chroma import Chroma

from peft import (
    PeftModel,
    LoraConfig,
    prepare_model_for_kbit_training,
    set_peft_model_state_dict,
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
torch.cuda.empty_cache()

@dataclass
class Metadata:
    """Metadata for processed documents"""
    filename: str
    chunk_count: int
    total_tokens: int
    processing_time: float

class Processor:
    def __init__(self):
        self.supported_extensions = {
            'pdf': self.process_pdf,
            'docx': self.process_docx,
            'txt': self.process_text,
            'md': self.process_text,
            'csv': self.process_text
        }

    def process_docx(self, document) -> str:
        try: 
            doc = docx.Document(document)
            content = []
            for section in doc.sections:
                header = section.header
                if header: 
                    content.append(paragraph.text for paragraph in header.paragraphs)
            
            content.append(paragraph.text for paragraph in doc.paragraphs)

            for table in doc.tables: 
                for row in table.rows: 
                    content.append(cell.text for cell in row.cells)
            
            return '\n'.join(filter(None, content))
        
        except Exception as e: 
            logger.error(f"Error processing DOCX: {str(e)}")
            raise ValueError(f"Failed to process DOCX: {str(e)}")

    def process_pdf(self, document) -> str:
        try: 
            pdf_reader = PyPDF2.PdfReader(document)
            return '\n'.join(
                page.extract_text().strip() for page in pdf_reader.pages if page.extract_text().strip()
            )        
        except Exception as e: 
            logger.error(f"Error processing PDF: {str(e)}")
            raise ValueError(f"Failed to process PDF: {str(e)}")

    def process_text(self, document):
        try:
            with open(document, "rb") as f:
                data = f.read()
            result = chardet.detect(data)
            encodings = [result['encoding'], 'utf-8', 'latin-1', 'ascii']
                
            for encoding in encodings:
                try:
                    if encoding:
                        return data.decode(encoding)
                except UnicodeDecodeError:
                    continue
                
            raise ValueError("Unable to decode file with any supported encoding")
        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}")
            raise ValueError(f"Failed to process text file: {str(e)}")
        
    def process_file(self, document) -> Tuple[str, Metadata]:
        start_time = time.time()
        
        # Extract the file extension and name
        file_extension = os.path.splitext(document)[1].lower().lstrip('.')
        file_name = os.path.splitext(document)[0].split('/')[-1]

        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        processor = self.supported_extensions[file_extension]
        content = processor(document)
        
        # Calculate basic metrics
        metadata = Metadata(
            filename=file_name,
            chunk_count=len(content.split('\n')),
            total_tokens=len(content.split()),
            processing_time=time.time() - start_time
        )
        
        return content, metadata, file_name

class DeepSeekApplication:
    def __init__(
        self,
        client_id: int, 
        ori_model,
        lora_weights_path,
        lora_config_path,
        prompt_template,
        chunk_size: int = 1000, # Chunk size
        chunk_overlap: int = 0, # Chunk overlap
    ):
        self.client_id = client_id
        self.ori_model = ori_model
        self.lora_weights_path = lora_weights_path
        self.lora_config_path = lora_config_path
        self.prompt_template = prompt_template
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        self.init_model()
        self.doc_processor = Processor()
        self.document_store = None
        self.document_metadata = {}
        
        with open(self.lora_config_path + "/client_{}.pkl".format(self.client_id), "rb") as f:
            self.client = pickle.load(f)
        
        with open(self.lora_config_path + "/user_permission_resource.pkl", "rb") as f:
            user_permissions_resource = pickle.load(f)
        
        self.client.set_managers(user_permissions_resource)
    
    def init_model(self):
        config = AutoConfig.from_pretrained(self.ori_model, trust_remote_code=True)
        self.model = AutoModelForCausalLM.from_pretrained(
            self.ori_model,
            config=config,
            load_in_8bit=True,
            torch_dtype=torch.float16,
            device_map="auto",
        )

        self.text_splitter = CharacterTextSplitter(
            chunk_size=self.chunk_size, 
            chunk_overlap=self.chunk_overlap,
        )

        self.recursive_text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000, 
            chunk_overlap=500,
        )

        model_kwargs = {
            'device': device
        }
        # Ensure that there are no out of memory issues
        encode_kwargs = { 
            'normalize_embeddings': True, 
            'batch_size': 8
        }
        self.embeddings = HuggingFaceEmbeddings(
            model_name="BAAI/bge-small-en-v1.5", # Embedding is different from the original model due to efficient embedding usage
            model_kwargs=model_kwargs,
            encode_kwargs=encode_kwargs
        )

        # Loading the model
        self.model = prepare_model_for_kbit_training(self.model)
        lora_weights = torch.load(self.lora_weights_path) # Get the LoRA weights
        config_peft = LoraConfig.from_pretrained(self.lora_config_path)
        self.model = PeftModel(self.model, config_peft)
        set_peft_model_state_dict(self.model, lora_weights, "default")

        self.tokenizer = AutoTokenizer.from_pretrained(self.ori_model)    
        self.tokenizer.pad_token_id = (
            0
        )
        self.tokenizer.padding_side = "left"

        self.model.half()
        self.model.eval() # Set the model to evaluation mode 
    
    def retrieve_relevant_docs(self, question: str, top_k: int, sim_threshold: float) -> List[str]:
        if self.document_store is None: 
            raise ValueError("There are no documents uploaded.")
        
        try: 
            scores = self.document_store.similarity_search(
                query=question, 
                k=top_k
            )

            print(scores)

            # Splitting the documents recursively 
            text_splits = self.recursive_text_splitter.split_documents([scores[0]])
            # Making a vector representation of the documents using embeddings
            vectorstore = Chroma.from_documents(documents=text_splits, embedding=self.embeddings)
            # Getting the most relevant bits of the documents 
            self.results_with_scores = vectorstore.similarity_search_with_score(question, k=top_k)

            return scores

        except Exception as e: 
            logger.error(f"Error retrieving relevant documents: {str(e)}")
            raise
    
    def preprocess_file(self, document: str) -> str: 
        document = ' '.join(document.split())
        document = document.replace('\t', ' ').strip()
        return document
    
    def clean_metadata(self, metadata: dict) -> dict:
        return {
            k: str(v) if v is not None else ""  # or remove it with a conditional
            for k, v in metadata.items()
            if v is not None  # Optional: skip None entirely
        }

    
    def load_documents(self, documents: List[str], metadata: Optional[Dict[str, Metadata]] = None, custom_text: bool = False) -> None:
        try:
            doc_chunks = []
            self.documents = documents
            documents_array = []

            def loading_documents(documents: List, documents_array: List, dict: bool = False):
                # Getting the documents content into the Document Langchain object
                if dict: 
                    documents_array = (
                        Document(
                            page_content=doc["context"], 
                            metadata=self.clean_metadata(doc["metadata"])
                        )
                        for doc in documents
                    )
                else:
                    for idx, data in enumerate(metadata):
                        doc = documents[idx]
                        documents_array.append(Document(
                            page_content = doc,
                            metadata = {
                                "filename": metadata[data].filename,
                                "chunk_count": metadata[data].chunk_count,
                                "total_tokens": metadata[data].total_tokens,
                                "processing_time": metadata[data].processing_time,
                            }
                        ))
                    
                return tuple(documents_array)
            
            if documents:
                self.uploaded_doc_present = True
                self.documents_array = loading_documents(documents, documents_array, dict=False) # Adding additional documents to the chunks
            else: 
                self.uploaded_doc_present = False
                self.documents_array = loading_documents(self.client.get_documents(), documents_array, dict=True) # Adding the documents of the clients they have access to

            splitted_docs = self.text_splitter.split_documents(self.documents_array)

            if not splitted_docs:
                raise ValueError("No documents to index. Check the output of your document processing step.")

            self.document_store = FAISS.from_documents(splitted_docs, self.embeddings)
            
            logger.info(f"Successfully loaded {len(doc_chunks)} chunks from {len(documents)} documents")
            
        except Exception as e:
            logger.error(f"Error loading documents: {str(e)}")
            raise    

    def generate_response(
        self,
        query: str, # The question to be asked
        deepseek, # DeepSeekApplication instance for processing documents
        top_k: int, # Number of highest probability vocabulary tokens to keep for top-k filtering
        top_p: int, # The smallest set of most probable tokens with probabilities that add up to top_p or higher are kept for generation
        num_beams: int, # Number of beams for beam search
        max_new_tokens: int,
        similarity_threshold: float,
        temp: float,
        context: bool, # Boolean to check if there are documents for context
        max_context_length: int = 2000 # The maximum tokens for the input context 
    ):
        start_time = time.time()
        
        try:
            relevant_document = self.retrieve_relevant_docs(query, top_k, similarity_threshold)[0].page_content

            if self.uploaded_doc_present:
                retrieved_bits = [
                    text.page_content for text, _ in self.results_with_scores
                ]
                metadata = self.results_with_scores[0][0].metadata
            else: 
                # Lower scores is more similar
                retrieved_bits = [
                    text.page_content for text, score in self.results_with_scores if score <= 0.7
                ]
                metadata = self.results_with_scores[0][0].metadata

            combined_texts = ' '.join(retrieved_bits)

            # Truncate or assign the full text as needed
            combined_context = combined_texts[:max_context_length] + "..." if len(combined_texts) > max_context_length else combined_texts

            prompt = self.construct_prompt(query, combined_context)

            generation_config = GenerationConfig(
                temperature=temp,
                top_p=top_p,
                top_k=top_k,
                num_beams=num_beams
            )

            with torch.no_grad():
                generated_output = deepseek.model.generate(
                    prompt.to(device),
                    generation_config=generation_config,
                    do_sample=True,
                    max_new_tokens=5000,
                )
            input_length = prompt.shape[0]
            output = deepseek.tokenizer.batch_decode(generated_output[:, input_length:], skip_special_tokens=True)[0]

            answer = {
                'content': self.post_processing(output),
                'metadata': {
                    'text_metadata': metadata,
                    'processing_time': time.time() - start_time,
                    'context_length': len(combined_context),
                    'query_length': len(query)
                }
            }
            return answer, relevant_document
            
        except Exception as e:
            logger.error(f"Error generating response: {str(e)}")
            raise

    def post_processing(self, output: str) -> str:
        match = re.search(r"</think>\s*(.*)", output)
        if match:
            answer = match.group(1).strip()
            answer = re.sub(r"[\\]boxed\{(.*?)}", r"\1", answer)
            return answer.strip()

        # Fallback if pattern not found
        return output.strip()
    
    def test_generation(self, prompt: str, context: str, max_context_length: int, temp: int, top_p: int, top_k: int, num_beams: int, max_new_tokens: int) -> str:
        # Truncate context if it is too long
        combined_context = context[:max_context_length] + "..." if len(context) > max_context_length else context
        
        prompt = self.construct_prompt(prompt, combined_context)
        
        inputs = self.tokenizer(prompt, return_tensors="pt")
        input_ids = inputs["input_ids"].to(device)
        generation_config = GenerationConfig(
            temperature=temp,
            top_p=top_p,
            top_k=top_k,
            num_beams=num_beams
        )
        with torch.no_grad():
            generated_output = self.model.generate(
                input_ids=input_ids,
                generation_config=generation_config,
                return_dict_in_generate=True,
                output_scores=True,
                max_new_tokens=max_new_tokens,
            )
        s = generated_output.sequences[0]
        output = self.tokenizer.decode(s)

        return self.post_processing(output)


    def construct_prompt(self, query: str, context: str) -> str: 

        messages = [
            {"role": "system", "content": """You are an expert assistant designed to answer questions accurately, helpfully and concise.

            By the user, you are given an optional context document and a user question. If the context is useful, use it. If it is missing, unclear, or irrelevant, rely on your own knowledge to answer as clearly and informatively as possible.
            
            Instructions:
            - If the context is relevant and useful, base your answer on it.
            - If the context is insufficient or empty, answer using your own understanding and general knowledge.
            - Always respond in complete, well-structured short sentences. 
            - Do not explain steps or show reasoning unless explicitly asked.
            - Avoid unnecessary sentences or filler. Be direct and informative.
            - Do not mention the context’s quality (e.g., avoid saying "The context is insufficient").
            - Your goal is to provide the best possible answer regardless of context quality.""",},

            {"role": "user", "content": f"""
                Context (may be empty or partial):
                {context}

                Question:
                {query}
            """},
        ]

        tokenized_chat = self.tokenizer.apply_chat_template(messages, tokenize=True, add_generation_prompt=True, return_tensors="pt")
        return tokenized_chat
