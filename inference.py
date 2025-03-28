import torch
import gradio as gr
import fire
import logging
import docx
import PyPDF2
import chardet
import time
from langchain_community.vectorstores import FAISS   
from typing import Tuple, List, Optional, Dict

from utils.prompt_template import PromptHelper
from transformers import AutoTokenizer, AutoModelForCausalLM, AutoConfig, GenerationConfig

from peft import (
    PeftModel,
    LoraConfig,
    prepare_model_for_kbit_training,
    set_peft_model_state_dict,
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

class Metadata:
    """Metadata for processed documents"""
    filename: str
    chunk_count: int
    total_tokens: int
    processing_time: float

class Processor:
    def __init__(self):
        self.supported_extensions = {
            'pdf': self.process_pdf,
            'docx': self.process_docx,
            'txt': self.process_text,
            'md': self.process_text,
            'csv': self.process_text
        }

    def process_docx(self, document) -> str:
        try: 
            doc = docx.Document(document)
            content = []
            for section in doc.sections:
                header = section.header
                if header: 
                    content.append(paragraph.text for paragraph in header.paragraphs)
            
            content.append(paragraph.text for paragraph in doc.paragraphs)

            for table in doc.tables: 
                for row in table.rows: 
                    content.append(cell.text for cell in row.cells)
            
            return '\n'.join(filter(None, content))
        
        except Exception as e: 
            logger.error(f"Error processing DOCX: {str(e)}")
            raise ValueError(f"Failed to process DOCX: {str(e)}")

    def process_pdf(self, document) -> str:
        try: 
            pdf_reader = PyPDF2.PdfReader(document)
            return '\n'.join(
                page.extract_text().strip() for page in pdf_reader.pages if page.extract_text().strip()
            )        
        except Exception as e: 
            logger.error(f"Error processing PDF: {str(e)}")
            raise ValueError(f"Failed to process PDF: {str(e)}")

    def process_text(self, document):
        try: 
            data = document.getvalue()
            result = chardet.detect(data)
            encodings = [result['encoding'], 'utf-8', 'latin-1', 'ascii']
                
            for encoding in encodings:
                try:
                    if encoding:
                        return data.decode(encoding)
                except UnicodeDecodeError:
                    continue
                
                raise ValueError("Unable to decode file with any supported encoding")
        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}")
            raise ValueError(f"Failed to process text file: {str(e)}")
        
    def process_file(self, document) -> Tuple[str, Metadata]:
        start_time = time.time()
            
        filename = document.name
        file_ext = filename.split('.')[-1].lower()

        if file_ext not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        processor = self.supported_extensions[file_ext]
        content = processor(document)
        
        # Calculate basic metrics
        metadata = Metadata(
            filename=filename,
            chunk_count=len(content.split('\n')),
            total_tokens=len(content.split()),
            processing_time=time.time() - start_time
        )
        
        return content, metadata

class DeepSeekApplication:
    def __init__(
        self,
        ori_model: str = "deepseek-ai/DeepSeek-R1-Distill-Qwen-1.5B", # The original model 
        lora_weights_path: str = "FL_output/pytorch_model.bin", # Path to the weights after LoRA
        lora_config_path: str = "FL_output", # Path to the config.json file after LoRA
        prompt_template: str = 'utils/prompt_template.json', # Prompt template for LLM
    ):
        self.ori_model = ori_model
        self.lora_weights_path = lora_weights_path
        self.lora_config_path = lora_config_path
        self.prompt_template = prompt_template

        self.init_model()
        self.doc_processor = Processor()
        self.document_store = None
        self.document_metadata = {}
    
    def init_model(self):
        self.prompter = PromptHelper(self.prompt_template)
        config = AutoConfig.from_pretrained(self.ori_model, trust_remote_code=True)
        self.model = AutoModelForCausalLM.from_pretrained(
            self.ori_model,
            config=config,
            load_in_8bit=True,
            torch_dtype=torch.float16,
            device_map="auto",
        )

        # Loading the model
        self.model = prepare_model_for_kbit_training(self.model)
        lora_weights = torch.load(self.lora_weights_path) # Get the LoRA weights
        config_peft = LoraConfig.from_pretrained(self.lora_config_path)
        self.model = PeftModel(self.model, config_peft)
        set_peft_model_state_dict(self.model, lora_weights, "default")

        self.tokenizer = AutoTokenizer.from_pretrained(self.ori_model)    
        self.tokenizer.pad_token_id = (
            0
        )
        self.tokenizer.padding_side = "left"

        self.model.half()
        self.model.eval() # Set the model to evaluation mode 
    
    def retrieve_relevant_docs(self, question: str, top_k: int, sim_threshold: float) -> List[str]:
        if self.document_store is None: 
            raise ValueError("There are no documents uploaded.")
        
        try: 
            scores = self.document_store.similarity_search_with_score(
                query=question, 
                k=top_k
            )

            relevant_docs = [
                document.page_conent for document, score in scores if score >= sim_threshold
            ]

            return relevant_docs

        except Exception as e: 
            logger.error(f"Error retrieving relevant documents: {str(e)}")
            raise
    
    def preprocess_file(self, document: str) -> str: 
        document = ' '.join(document.split())
        document = document.replace('\t', ' ').strip()
        return document
    
    def load_documents(self, documents: List[str], metadata: Optional[Dict[str, Metadata]] = None) -> None:
        try:
            doc_chunks = []
            
            for doc in documents:
                cleaned_doc = self.preprocess_file(doc)
                if cleaned_doc:
                    chunks = self.text_splitter.split_text(cleaned_doc)
                    doc_chunks.extend(chunks)
            
            if not doc_chunks:
                raise ValueError("No valid document content found after processing.")
            
            self.vector_store = FAISS.from_texts(
                texts=doc_chunks,
                embedding=self.embeddings
            )
            
            if metadata:
                self.document_metadata.update(metadata)
            
            logger.info(f"Successfully loaded {len(doc_chunks)} chunks from {len(documents)} documents")
            
        except Exception as e:
            logger.error(f"Error loading documents: {str(e)}")
            raise    

    def generate_response(
        self,
        query: str,
        deepseek,
        top_k: int,
        top_p: int,
        num_beams: int,
        max_new_tokens: int,
        similarity_threshold: float,
        temp: float,
        context: bool = True,
        max_context_length: int = 2000
    ):
        start_time = time.time()
        
        try:
            if context:
                context = self.retrieve_relevant_docs(query, top_k, similarity_threshold)
                
                # Truncate context if too long
                combined_context = ' '.join(context)
                if len(combined_context) > max_context_length:
                    combined_context = combined_context[:max_context_length] + "..."
                
                prompt = self.construct_prompt(query, combined_context)
                inputs = deepseek.tokenizer(prompt, return_tensors="pt")
                input_ids = inputs["input_ids"].to(device)
                generation_config = GenerationConfig(
                    temperature=temp,
                    top_p=top_p,
                    top_k=top_k,
                    num_beams=num_beams
                )
                with torch.no_grad():
                    generated_output = deepseek.model.generate(
                        input_ids=input_ids,
                        generation_config=generation_config,
                        return_dict_in_generate=True,
                        output_scores=True,
                        max_new_tokens=max_new_tokens,
                    )
                s = generated_output.sequences[0]
                output = deepseek.tokenizer.decode(s)
                response = self.prompter.get_response(prompt)
                answer = {
                    'content': response,
                    'metadata': {
                        'processing_time': time.time() - start_time,
                        'context_length': len(combined_context),
                        'query_length': len(query)
                    }
                }
                yield self.prompter.get_response(output)
            else: 
                # If there is no context construct a "normal" prompt
                prompt = self.prompter.generate_prompt(query, "")
                inputs = deepseek.tokenizer(prompt, return_tensors="pt")
                input_ids = inputs["input_ids"].to(device)
                generation_config = GenerationConfig(
                    temperature=temp,
                    top_p=top_p,
                    top_k=top_k,
                    num_beams=num_beams
                )
                with torch.no_grad():
                    generated_output = deepseek.model.generate(
                        input_ids=input_ids,
                        generation_config=generation_config,
                        return_dict_in_generate=True,
                        output_scores=True,
                        max_new_tokens=max_new_tokens,
                    )
                s = generated_output.sequences[0]
                output = deepseek.tokenizer.decode(s)
                return self.prompter.get_response(output), "test"
            
        except Exception as e:
            logger.error(f"Error generating response: {str(e)}")
            raise

    def construct_prompt(self, query: str, context: str) -> str:
        """Construct an enhanced prompt template"""
        return f"""
        Context Information:
        {context}

        Question: {query}

        Please provide a comprehensive answer based on the context above. Consider:
        1. Direct relevance to the question
        2. Accuracy of information
        3. Completeness of response
        4. Clarity and coherence

        Answer:
        """

def run(
    ori_model: str = "deepseek-ai/DeepSeek-R1-Distill-Qwen-1.5B", # The original model 
    lora_weights_path: str = "FL_output/pytorch_model.bin", # Path to the weights after LoRA
    lora_config_path: str = "FL_output", # Path to the config.json file after LoRA
    prompt_template: str = 'utils/prompt_template.json', # Prompt template for LLM
):
    deepseek = DeepSeekApplication()

    def evaluate(
        question: str, # The question to be asked
        uploaded_documents: str = None, # The corresponding document(s)
	    custom_text: str = None,
        temp: float = 0.1, # Temperature to module the next token probabilities
        top_p: float = 0.75, # Only the smallest set of the most probable tokens with probabilities that add up to top_p or higher are kept for generation
        top_k: int = 40, # Number of highest probability vocabulary tokens to keep for top-k-filtering
        num_beams: int = 4, # Number of beams for beam search
        max_new_tokens: int = 128,
        **kwargs,
    ):  
        documents = []
        metadata = {}
        if uploaded_documents['files']: 
            for file in uploaded_documents: 
                content, metadata_doc = deepseek.doc_processor.process_file(file)
                documents.append(content)
                metadata[file.name] = metadata_doc

            deepseek.load_documents(documents, metadata)
            response = deepseek.generate_response(question, deepseek, top_k, top_p, num_beams, max_new_tokens, 0.0, temp)
        else:
            response = deepseek.generate_response(question, deepseek, top_k, top_p, num_beams, max_new_tokens, 0.0, temp, context=False)
        return response, metadata

    UI = gr.Interface(
        fn=evaluate,
        inputs=[
            gr.components.Textbox(
                lines=2,
                label="❓Question",
                info="Upload documents below to ask questions about the content.",
            ),
            gr.MultimodalTextbox(
                file_count='multiple',
                placeholder="Upload your documents here.",
                label="📁 Document Input",
                show_label=True,
                info="Supported formats: PDF, DOCX, TXT"
            ),
            gr.components.Textbox(
                lines=1,
                label="📃 Or paste text",
                info="Enter text directly. Each paragraph will be processed seperately."
            ),
            gr.components.Slider(
                minimum=0, maximum=1, value=0.6, label="🌡️ Temperature"
            ),
            gr.components.Slider(
                minimum=0, maximum=1, value=0.75, label="Top p"
            ),
            gr.components.Slider(
                minimum=0, maximum=100, step=1, value=40, label="Top k"
            ),
            gr.components.Slider(
                minimum=1, maximum=4, step=1, value=4, label="Beams"
            ),
            gr.components.Slider(
                minimum=1, maximum=2000, step=1, value=500, label="Max tokens"
            ),
        ],
        outputs=[
            gr.Textbox(
                lines=10,
                label="🔮 Output",
                info="Output of the DeepSeek model."
            ),
            gr.Textbox(
                lines=10, 
                label="📊 Document Info",
                info="Meta Data of the input documents."
            )
        ],
        title="🔎 DeepSeek Q&A",
        description=""" 
            ### Document Analysis and Question Answering.
            # Upload documents or paste text to ask questions about the content.
        """ ,
        theme=gr.themes.Default(primary_hue=gr.themes.colors.blue, secondary_hue=gr.themes.colors.blue),
        submit_btn="Generate Response",
        flagging_mode="never"
    ).queue()

    UI.launch(share=True, server_port=7860)

if __name__ == "__main__":
    fire.Fire(run)
